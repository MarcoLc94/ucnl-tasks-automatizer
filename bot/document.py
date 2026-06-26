from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(
    course_name: str,
    task_title: str,
    response_text: str,
    student_name: str,
    output_dir: Path,
) -> Path:
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Encabezado institucional
    header = doc.add_paragraph("Universidad Ciudadana de Nuevo León")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.runs[0]
    run.bold = True
    run.font.size = Pt(13)

    sub = doc.add_paragraph("Ingeniería en Desarrollo de Software")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Materia y título
    materia = doc.add_paragraph()
    materia.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = materia.add_run(course_name)
    r.bold = True
    r.font.size = Pt(12)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = titulo.add_run(task_title)
    r2.bold = True
    r2.font.size = Pt(12)

    doc.add_paragraph()

    # Datos del alumno
    alumno = doc.add_paragraph()
    alumno.alignment = WD_ALIGN_PARAGRAPH.CENTER
    alumno.add_run(f"Alumno: {student_name}").font.size = Pt(11)

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}").font.size = Pt(11)

    doc.add_paragraph()

    # Contenido
    for paragraph in response_text.split("\n"):
        p = doc.add_paragraph(paragraph)
        p.runs[0].font.size = Pt(11) if p.runs else None

    # Guardar
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_title = "".join(c for c in task_title if c.isalnum() or c in " _-")[:50].strip()
    path = output_dir / f"{safe_title}.docx"
    doc.save(str(path))
    return path
